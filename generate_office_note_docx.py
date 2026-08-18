# -*- coding: utf-8 -*-
"""
generate_office_note_docx.py
============================
Generates a Word (.docx) document for the Hindi Office Note:
  "MB Withdrawal Alert — Package DD2/BAD/PR/25-26/6"
  dated 17 August 2026

All body text uses Mangal font (standard for Hindi/Devanagari in
Indian government documents).

Requirements:
    pip install python-docx

Run:
    python generate_office_note_docx.py
Output:
    office_note_MB_withdrawal_17Aug2026.docx  (same folder)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── constants ───────────────────────────────────────────────────────────────
MANGAL       = "Mangal"
ARIAL_UNICODE = "Arial Unicode MS"   # fallback if Mangal unavailable
BLUE         = RGBColor(0x00, 0x35, 0x80)
RED          = RGBColor(0xCC, 0x00, 0x00)
BLACK        = RGBColor(0x1A, 0x1A, 0x1A)
LIGHT_BLUE   = RGBColor(0xE8, 0xEE, 0xF8)
LIGHT_RED    = RGBColor(0xFF, 0xF5, 0xF5)

OUTPUT_FILE  = os.path.join(os.path.dirname(__file__),
                            "office_note_MB_withdrawal_17Aug2026.docx")


# ── helpers ──────────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=13, color=BLACK, font_name=MANGAL):
    run.font.name        = font_name
    run.font.size        = Pt(size)
    run.font.bold        = bold
    run.font.color.rgb   = color
    # ensure Devanagari rendering in Word
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:hint'), 'cs')
    rFonts.set(qn('w:cs'),   font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text="", bold=False, size=13, color=BLACK,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
                  highlight_words=None):
    """Add a paragraph. highlight_words = list of (word, color) tuples."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    if highlight_words:
        # split text around highlighted phrases
        remaining = text
        for phrase, ph_color in highlight_words:
            parts = remaining.split(phrase, 1)
            if len(parts) == 2:
                if parts[0]:
                    r = p.add_run(parts[0])
                    set_font(r, bold=bold, size=size, color=color)
                r2 = p.add_run(phrase)
                set_font(r2, bold=True, size=size, color=ph_color)
                remaining = parts[1]
        if remaining:
            r = p.add_run(remaining)
            set_font(r, bold=bold, size=size, color=color)
    else:
        r = p.add_run(text)
        set_font(r, bold=bold, size=size, color=color)
    return p


def add_heading(doc, text, level_color=BLUE, size=13, bg=LIGHT_BLUE):
    """Section heading with left blue border (simulated via shading + bold)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    # shading
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '{:02X}{:02X}{:02X}'.format(*bg))
    pPr.append(shd)
    # left border
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*level_color))
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_font(run, bold=True, size=size, color=level_color)
    return p


def add_divider(doc, thick=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12' if thick else '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '003580')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def shade_paragraph(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)


def add_bullet(doc, text, indent=0.5, bold_part=None, bold_color=BLACK):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.line_spacing  = Pt(20)
    if bold_part and bold_part in text:
        parts = text.split(bold_part, 1)
        r1 = p.add_run(parts[0])
        set_font(r1, size=12)
        r2 = p.add_run(bold_part)
        set_font(r2, bold=True, size=12, color=bold_color)
        r3 = p.add_run(parts[1])
        set_font(r3, size=12)
    else:
        r = p.add_run(text)
        set_font(r, size=12)
    return p


# ── main ─────────────────────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # ── page margins (A4, Office-style) ──────────────────────────────────────
    section = doc.sections[0]
    section.page_height  = Cm(29.7)
    section.page_width   = Cm(21.0)
    section.left_margin  = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin= Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    shade_paragraph(title_p, '003580')
    r = title_p.add_run("कार्यालय टिप्पणी  |  Office Note")
    set_font(r, bold=True, size=16, color=RGBColor(0xFF,0xFF,0xFF))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(8)
    shade_paragraph(sub_p, '003580')
    r2 = sub_p.add_run("राजस्थान लोक निर्माण विभाग  |  Rajasthan Public Works Department")
    set_font(r2, bold=False, size=11, color=RGBColor(0xCC,0xDD,0xFF))

    add_divider(doc, thick=True)

    # ══════════════════════════════════════════════════════════════════════════
    # META TABLE
    # ══════════════════════════════════════════════════════════════════════════
    meta = [
        ("विषय",          "प्रथम चालू बिल — स्थायी पुनर्स्थापना कार्य"),
        ("पैकेज क्रमांक", "DD2/BAD/PR/25-26/6"),
        ("ठेकेदार",        "मेसर्स एस.के. कंस्ट्रक्शन"),
        ("दिनांक",         "17 अगस्त 2026"),
        ("टिप्पणी दिनांक", "18 अगस्त 2026"),
        ("संदर्भ",          "IFMS WAM Abstract Submission एवं Electronic MB का अनधिकृत Withdrawal"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment    = WD_TABLE_ALIGNMENT.LEFT
    tbl.style        = 'Table Grid'
    for i, (label, value) in enumerate(meta):
        row = tbl.rows[i]
        # label cell
        lc = row.cells[0]
        lc.width = Cm(5.5)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_after = Pt(3)
        lr = lp.add_run(label)
        set_font(lr, bold=True, size=12, color=BLUE)
        # shading alternating rows
        fill = 'F0F4FF' if i % 2 == 0 else 'FFFFFF'
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd   = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tc_pr.append(shd)
        # value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_after = Pt(3)
        vr = vp.add_run(value)
        bold_val = label in ("पैकेज क्रमांक", "दिनांक")
        set_font(vr, bold=bold_val, size=12, color=BLACK)

    doc.add_paragraph()  # spacer

    add_divider(doc, thick=True)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — तथ्यात्मक विवरण
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "तथ्यात्मक विवरण")

    add_paragraph(doc,
        "दिनांक 17 अगस्त 2026 को अपरान्ह 3:00 बजे से 4:30 बजे के मध्य उपर्युक्त कार्य की "
        "5 सड़कों के प्रथम चालू बिल के Abstract, IFMS WAM पोर्टल पर Submit किए गए थे।",
        highlight_words=[
            ("17 अगस्त 2026", BLUE),
            ("अपरान्ह 3:00 बजे से 4:30 बजे", BLUE),
            ("5 सड़कों के प्रथम चालू बिल के Abstract", BLUE),
        ])

    add_paragraph(doc,
        "उक्त Abstract को आपसे OTP प्राप्त कर Approve किया गया तथतपश्चात् WAM पर "
        "Bill Initiate किया गया, जिसका प्रिंटआउट संलग्न है।",
        highlight_words=[
            ("आपसे OTP प्राप्त कर Approve किया गया", BLUE),
            ("Bill Initiate", BLUE),
        ])

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — आपत्तिजनक घटना
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "आपत्तिजनक घटना")

    add_paragraph(doc,
        "सायं 4:30 बजे के लगभग उपर्युक्त समस्त कार्यों की Measurement Books (MB), "
        "IFMS Electronic MB रिकॉर्ड से Withdrawn पाई गईं।",
        highlight_words=[
            ("सायं 4:30 बजे के लगभग", RED),
            ("Withdrawn पाई गईं।", RED),
        ])

    add_paragraph(doc,
        "यह घटना तब ही संभव है जबकि उन्हें पोर्टल https://mb.rajasthan.gov.in/ "
        "— जो IFMS की इलेक्ट्रॉनिक MB साइट है — से Withdraw किया गया हो।")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — संज्ञान हेतु अनुरोध
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "संज्ञान हेतु अनुरोध")

    add_paragraph(doc,
        "कृपया इस विषय का गंभीरता से संज्ञान लें और निम्नलिखित बिन्दुओं पर विचार करें:")

    # numbered points
    numbered = [
        ("यदि यह कार्य आपकी पूर्वानुमति के बिना किया गया है",
         " — तो यह Digital Signature का दुरुपयोग / अनधिकृत उपयोग है, "
         "जो गंभीर वित्तीय एवं प्रशासनिक अनियमितता की श्रेणी में आता है।"),
        ("पारदर्शिता एवं सावधानी",
         " के दृष्टिगत यह विषय आपके ध्यान में लाया जा रहा है।"),
        ("यदि MB Withdrawal आपकी जानकारी में हुई है",
         " — तो कृपया उसका कारण एवं औचित्य स्पष्ट करें ताकि अभिलेख "
         "उचित रूप से अद्यतन किए जा सकें।"),
    ]
    for bold_part, rest in numbered:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(21)
        r1 = p.add_run(bold_part)
        set_font(r1, bold=True, size=12, color=BLUE)
        r2 = p.add_run(rest)
        set_font(r2, size=12, color=BLACK)

    # sub-bullets under point 2
    bullets_2 = [
        "Digital Signature का किसी भी प्रकार का दुरुपयोग रोका जा सके।",
        "समस्त Transactions की सुरक्षा एवं पारदर्शिता सुनिश्चित की जा सके।",
        "भविष्य में ऐसी स्थिति पुनः उत्पन्न न हो।",
    ]
    for b in bullets_2:
        add_bullet(doc, b, indent=1.5)

    # ══════════════════════════════════════════════════════════════════════════
    # FILING REQUEST — highlighted box
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()

    # box heading
    box_head = doc.add_paragraph()
    box_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    box_head.paragraph_format.space_after  = Pt(4)
    box_head.paragraph_format.space_before = Pt(8)
    shade_paragraph(box_head, 'CC0000')
    rh = box_head.add_run("  फाइल करने का अनुरोध")
    set_font(rh, bold=True, size=13, color=RGBColor(0xFF,0xFF,0xFF))

    # box body paragraphs
    box_lines = [
        "यदि MB Withdrawal आपकी जानकारी में हुई है, तो इस नोट को उचित अभिलेख / फाइल में "
        "दर्ज (File) किए जाने का अनुरोध है —",
    ]
    for line in box_lines:
        bp = doc.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.left_indent   = Cm(0.3)
        bp.paragraph_format.space_after   = Pt(4)
        bp.paragraph_format.line_spacing  = Pt(21)
        shade_paragraph(bp, 'FFF5F5')
        r1 = bp.add_run("यदि MB Withdrawal आपकी जानकारी में हुई है")
        set_font(r1, bold=True, size=12, color=RED)
        r2 = bp.add_run(
            ", तो इस नोट को उचित अभिलेख / फाइल में दर्ज (File) किए जाने का अनुरोध है —")
        set_font(r2, bold=True, size=12, color=BLACK)

    box_bullets = [
        "इस कार्यालय टिप्पणी को संबंधित कार्य की फाइल में संलग्न किया जाए।",
        "आपकी स्वीकृति / टिप्पणी / कारण अभिलेख पर अंकित किया जाए।",
        "आवश्यकतानुसार उच्च अधिकारियों को भी अवगत कराया जाए।",
    ]
    for bb in box_bullets:
        bp2 = doc.add_paragraph()
        bp2.paragraph_format.left_indent  = Cm(1.0)
        bp2.paragraph_format.space_after  = Pt(3)
        bp2.paragraph_format.line_spacing = Pt(20)
        shade_paragraph(bp2, 'FFF5F5')
        r = bp2.add_run("• " + bb)
        set_font(r, size=12, color=BLACK)

    closing_p = doc.add_paragraph()
    closing_p.paragraph_format.left_indent  = Cm(0.3)
    closing_p.paragraph_format.space_after  = Pt(8)
    closing_p.paragraph_format.line_spacing = Pt(21)
    shade_paragraph(closing_p, 'FFF5F5')
    rc = closing_p.add_run(
        "यह अनुरोध विभागीय पारदर्शिता, डिजिटल सुरक्षा एवं लोकहित में किया जा रहा है।")
    set_font(rc, size=12, color=BLACK)

    # ══════════════════════════════════════════════════════════════════════════
    # ANNEXURE
    # ══════════════════════════════════════════════════════════════════════════
    ann_p = doc.add_paragraph()
    ann_p.paragraph_format.space_after  = Pt(6)
    ann_p.paragraph_format.line_spacing = Pt(20)
    shade_paragraph(ann_p, 'F9F9E8')
    ra1 = ann_p.add_run("संलग्नक: ")
    set_font(ra1, bold=True, size=12, color=RGBColor(0x88, 0x77, 0x00))
    ra2 = ann_p.add_run(
        "WAM Bill Initiation का प्रिंटआउट (5 सड़कें, Package No. DD2/BAD/PR/25-26/6)")
    set_font(ra2, size=12, color=BLACK)

    # ══════════════════════════════════════════════════════════════════════════
    # GENERAL NOTE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "निवेदन")
    add_paragraph(doc,
        "यह टिप्पणी विभागीय हित, डिजिटल लेनदेन की सुरक्षा, एवं सार्वजनिक धन की "
        "पारदर्शिता के उद्देश्य से प्रस्तुत की जा रही है।")

    add_divider(doc, thick=True)

    # ══════════════════════════════════════════════════════════════════════════
    # SIGNATURE BLOCK
    # ══════════════════════════════════════════════════════════════════════════
    sig_tbl = doc.add_table(rows=2, cols=2)
    sig_tbl.style = 'Table Grid'
    sig_data = [
        [("प्रस्तुतकर्ता का नाम:", ""), ("पदनाम:", "")],
        [("हस्ताक्षर:", ""), ("दिनांक:", "18 अगस्त 2026")],
    ]
    for row_idx, row_data in enumerate(sig_data):
        for col_idx, (label, value) in enumerate(row_data):
            cell = sig_tbl.rows[row_idx].cells[col_idx]
            cp   = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(16)
            rl = cp.add_run(label + "  ")
            set_font(rl, bold=True, size=12, color=BLUE)
            rv = cp.add_run(value if value else "_" * 30)
            set_font(rv, size=12, color=BLACK)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # FOOTER NOTE
    # ══════════════════════════════════════════════════════════════════════════
    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_p.paragraph_format.space_before = Pt(10)
    rf = foot_p.add_run(
        "यह दस्तावेज़ Legal Correspondence Manager के अंतर्गत तैयार किया गया है  |  "
        "github.com/CRAJKUMARSINGH/Legal-Correspondence_Manager")
    set_font(rf, size=9, color=RGBColor(0x77, 0x77, 0x77))

    # ── save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"[OK] Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
